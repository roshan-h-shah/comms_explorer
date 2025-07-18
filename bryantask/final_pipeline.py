# Straight Markdown + SQL Calling COmbines
# final_pipeline.py

# final_pipeline.py
'''
import os
import logging
import duckdb
import broadsql
from datacenter import run_scrape_and_markdown
from ooni import scrape_ooni_explorer
from cloudflare import fetch_and_format_markdown
from country_code_converter import get_alpha2_from_country_name
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.prompts import PromptTemplate
import asyncio

from docx import Document
from docx.shared import Pt
from docx2pdf import convert       # pip install python-docx docx2pdf
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas  # pip install reportlab

# — Init —
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Connect to DuckDB database
con = duckdb.connect("bryan.db")

# LLM Setup
llm = ChatGoogleGenerativeAI(
    model="gemini-2.5-flash",
    google_api_key=GOOGLE_API_KEY,
    temperature=0
)

def query_llm(agent_input: str) -> str:
    """Basic LLM wrapper."""
    response = llm.invoke(agent_input)
    logger.info(f"LLM Response:\n{response.content.strip()}")
    return response.content.strip()

def combined_pipeline(
    user_query: str,
    sql_tables: list[str],
    test_names: list[str],    # e.g. ['signal','whatsapp']
    only: str = "",           # 'none' or 'anomalies'
    horizon: int = 30         # days back for OONI & Radar
) -> str:
    """
    1) SQL-RAG to pick countries_list.
    2) Scrape data centers.
    3) Scrape OONI Explorer per country/test.
    4) Fetch Cloudflare Radar per country.
    5) LLM prompt combining all four contexts.
    """

    # --- 1) Build SQL context ---
    sql_blocks = []
    countries_list: list[str] = []
    for i, table in enumerate(sql_tables):
        logger.info(f"[SQL] Querying table: {table}")
        df = con.execute(f"SELECT * FROM {table}").df()
        if i == 0:
            countries_list = broadsql.extract_relevant_rows(df, user_query)
        filtered = broadsql.filter_df(df, "Country", countries_list)
        md = broadsql.df_to_markdown(filtered)
        sql_blocks.append(f"### Table: {table}\n{md}")
    sql_context = "\n\n".join(sql_blocks) or "No SQL data found."

    # --- 2) Data-Center context ---
    logger.info(f"[DC] Scraping data centers for: {countries_list}")
    dc_md = run_scrape_and_markdown(countries_list)
    dc_context = f"### Data Centers\n{dc_md}"

    # --- 3) OONI Explorer context ---
    logger.info(f"[OONI] Scraping tests {test_names} for: {countries_list}")
    ooni_blocks = []
    for test_name in test_names:
        for country in countries_list:
            alpha2 = get_alpha2_from_country_name(country) or ""
            md_table, anomaly_cnt, accessible_cnt = asyncio.run(
                scrape_ooni_explorer(
                    test_name=test_name,
                    country=alpha2,
                    only=only,
                    horizon=horizon
                )
            )
            ooni_blocks.append(
                f"#### {test_name.capitalize()} results for {country}\n"
                f"{md_table}\n"
                f"- **Anomalies**: {anomaly_cnt}   **Accessible**: {accessible_cnt}\n"
            )
    ooni_context = "\n\n".join(ooni_blocks) or "No OONI data found."

    # --- 4) Cloudflare Radar context ---
    logger.info(f"[CF] Fetching Radar data for: {countries_list}")
    radar_blocks = []
    date_range = f"{horizon}d"
    for country in countries_list:
        alpha2 = get_alpha2_from_country_name(country) or ""
        if not alpha2:
            continue
        # returns full Markdown for that one country
        radar_md = fetch_and_format_markdown(country=alpha2, date_range=date_range)
    
        radar_blocks.append(
            f"#### Cloudflare Radar ({date_range}) for {country}\n\n{radar_md}"
        )
     
    radar_context = "\n\n".join(radar_blocks) or "No Cloudflare Radar data found."
    logger.info(f"Radar Context for LLM: {radar_context}")
    # --- 5) Final LLM prompt ---
    prompt = PromptTemplate.from_template(
        """
You are an expert data analyst and technical writer. You have four blocks of data:

1) **SQL Tables**  
{sql_context}

2) **Web-scraped Data Centers**  
{dc_context}

3) **OONI Explorer Results**  
{ooni_context}

4) **Cloudflare Radar Summary**  
{radar_context}

User’s question:
> {user_query}

Produce a polished Markdown report with:
- **Executive Summary** (key findings)
- **Section for SQL Data** (tables/bullets)
- **Section for Data Centers** (bullets or table)
- **Section for OONI Explorer** (raw tables + summary counts)
- **Section for Cloudflare Radar** (tables per metric)
- **Overall Insights**

Use Markdown headings, bullet lists, and tables throughout.
"""
    ).format(
        sql_context=sql_context,
        dc_context=dc_md,
        ooni_context=ooni_context,
        radar_context=radar_context,
        user_query=user_query
    )

    logger.info("[LLM] Generating final report")
    return query_llm(prompt)

# --- Export helpers unchanged (markdown_to_docx, md_to_pdf, docx_to_pdf) ---

def markdown_to_docx(md_text: str, docx_path: str = "report.docx") -> str:
    """Convert markdown to Word with proper headings & bullets."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for line in md_text.splitlines():
        if line.startswith("### "):
            h = doc.add_heading(level=3); h.add_run(line[4:])
        elif line.startswith("## "):
            h = doc.add_heading(level=2); h.add_run(line[3:])
        elif line.startswith("# "):
            h = doc.add_heading(level=1); h.add_run(line[2:])
        elif line.startswith("* "):
            p = doc.add_paragraph(style='List Bullet'); p.add_run(line[2:])
        else:
            doc.add_paragraph(line)
    doc.save(docx_path)
    logger.info(f"Wrote Word doc: {docx_path}")
    return docx_path



if __name__ == "__main__":
    tables = ["mcc_mnc_table", "traforama_isp_list", "mideye_mobile_network_list"]
    report = combined_pipeline(
        user_query="Describe ISPs and signal/whatsapp anomalies in South Asia, and include Radar metrics.",
        sql_tables=tables,
        test_names=["signal", "whatsapp"],
        only="anomalies",
        horizon=30
    )
    print(report) 
'''

# final_pipeline.py
'''
import os
import logging
import duckdb
import broadsql
from datacenter import run_scrape_and_markdown
from ooni import scrape_ooni_explorer
from cloudflare import fetch_and_format_markdown
from country_code_converter import get_alpha2_from_country_name
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.prompts import PromptTemplate
import asyncio
import concurrent.futures

from docx import Document
from docx.shared import Pt
from docx2pdf import convert       # pip install python-docx docx2pdf
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas  # pip install reportlab

# — Init —
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Connect to DuckDB database
DB_PATH = os.path.join(os.path.dirname(__file__), "bryan.db")
con = duckdb.connect(DB_PATH)

# LLM Setup
llm = ChatGoogleGenerativeAI(
    model="gemini-2.5-flash",
    google_api_key=GOOGLE_API_KEY,
    temperature=0
)

def query_llm(agent_input: str) -> str:
    """Basic LLM wrapper."""
    response = llm.invoke(agent_input)
    logger.info(f"LLM Response:\n{response.content.strip()}")
    return response.content.strip()

# ----------------------------------------
# Section‐specific LLM callers
# ----------------------------------------
def answer_sql_section(user_query: str, sql_context: str) -> str:
    """Summarize the SQL‐extracted data for the user's question."""
    prompt = PromptTemplate.from_template(
        """
You are a data analyst. Given these SQL‐extracted tables:

{sql_context}

Answer the user’s question:
"{user_query}"

Focus only on the SQL data. Provide a concise Markdown bullet list of the relevant findings.
"""
    ).format(sql_context=sql_context, user_query=user_query)
    return query_llm(prompt)

def answer_dc_section(dc_context: str) -> str:
    """Summarize data‐center info."""
    prompt = PromptTemplate.from_template(
        """
You are a network infrastructure expert. Given these data‐center entries:

{dc_context}

List each data center with its name, type, and address as a Markdown table.
"""
    ).format(dc_context=dc_context)
    return query_llm(prompt)

def answer_ooni_section(ooni_context: str) -> str:
    """Summarize OONI test results."""
    prompt = PromptTemplate.from_template(
        """
You are a network measurement specialist. Here are OONI Explorer results:

{ooni_context}

For each country and test, list anomalies vs accessible counts in bullet points.
"""
    ).format(ooni_context=ooni_context)
    return query_llm(prompt)

def answer_radar_section(radar_context: str) -> str:
    """Summarize Cloudflare Radar metrics."""
    prompt = PromptTemplate.from_template(
        """
You are a traffic analysis expert. Here is Cloudflare Radar data:

{radar_context}

For each metric, produce a Markdown table showing the top categories and their percentages.
"""
    ).format(radar_context=radar_context)
    return query_llm(prompt)

# ----------------------------------------
# Combined pipeline
# ----------------------------------------
def combined_pipeline(
    user_query: str,
    sql_tables: list[str],
    test_names: list[str],    # e.g. ['signal','whatsapp']
    only: str = "",           # 'none' or 'anomalies'
    horizon: int = 30         # days back for OONI & Radar
) -> str:
    """
    1) Runs SQL‐RAG to pick countries_list.
    2) Scrapes data centers.
    3) Scrapes OONI Explorer per country/test.
    4) Fetches Cloudflare Radar per country.
    5) Calls 4 smaller LLM prompts in parallel and stitches results.
    """

    # --- 1) Build SQL context ---
    sql_blocks = []
    countries_list: list[str] = []
    for i, table in enumerate(sql_tables):
        logger.info(f"[SQL] Querying table: {table}")
        df = con.execute(f"SELECT * FROM {table}").df()
        if i == 0:
            countries_list = broadsql.extract_relevant_rows(df, user_query)
        filtered = broadsql.filter_df(df, "Country", countries_list)
        sql_blocks.append(f"### {table}\n{broadsql.df_to_markdown(filtered)}")
    sql_context = "\n\n".join(sql_blocks) or "No SQL data found."

    # --- 2) Data‐Center context ---
    logger.info(f"[DC] Scraping data centers for: {countries_list}")
    dc_md = run_scrape_and_markdown(countries_list)
    dc_context = dc_md  # pure markdown table

    # --- 3) OONI Explorer context ---
    logger.info(f"[OONI] Scraping tests {test_names} for: {countries_list}")
    ooni_blocks = []
    for test_name in test_names:
        for country in countries_list:
            alpha2 = get_alpha2_from_country_name(country) or ""
            md_table, anomaly_cnt, accessible_cnt = asyncio.run(
                scrape_ooni_explorer(
                    test_name=test_name,
                    country=alpha2,
                    only=only,
                    horizon=horizon
                )
            )
            ooni_blocks.append(
                f"#### {test_name.title()} – {country}\n"
                f"{md_table}\n"
                f"- Anomalies: {anomaly_cnt} &nbsp; Accessible: {accessible_cnt}\n"
            )
    ooni_context = "\n\n".join(ooni_blocks) or "No OONI data found."

    # --- 4) Cloudflare Radar context ---
    logger.info(f"[CF] Fetching Radar data for: {countries_list}")
    radar_blocks = []
    date_range = f"{horizon}d"
    for country in countries_list:
        alpha2 = get_alpha2_from_country_name(country) or ""
        radar_md = fetch_and_format_markdown(country=alpha2, date_range=date_range)
        radar_blocks.append(f"#### {country} (Radar {date_range})\n{radar_md}")
    radar_context = "\n\n".join(radar_blocks) or "No Radar data found."

    # ----------------------------------------
    # 5) Parallel LLM calls for each section
    # ----------------------------------------
    with concurrent.futures.ThreadPoolExecutor(max_workers=4) as executor:
        future_sql   = executor.submit(answer_sql_section,   user_query, sql_context)
        future_dc    = executor.submit(answer_dc_section,    dc_context)
        future_ooni  = executor.submit(answer_ooni_section,  ooni_context)
        future_radar = executor.submit(answer_radar_section, radar_context)

        sql_ans   = future_sql.result()
        dc_ans    = future_dc.result()
        ooni_ans  = future_ooni.result()
        radar_ans = future_radar.result()

    # --- Stitch final report ---
    report = "\n\n".join([
        "## SQL Data",
        sql_ans,
        "## Data Centers",
        dc_ans,
        "## OONI Explorer Results",
        ooni_ans,
        "## Cloudflare Radar",
        radar_ans
    ])
    return report




def markdown_to_docx(md_text: str, docx_path: str = "report.docx") -> str:
    """Convert markdown to Word with proper headings & bullets."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for line in md_text.splitlines():
        if line.startswith("### "):
            h = doc.add_heading(level=3); h.add_run(line[4:])
        elif line.startswith("## "):
            h = doc.add_heading(level=2); h.add_run(line[3:])
        elif line.startswith("# "):
            h = doc.add_heading(level=1); h.add_run(line[2:])
        elif line.startswith("* "):
            p = doc.add_paragraph(style='List Bullet'); p.add_run(line[2:])
        else:
            doc.add_paragraph(line)
    doc.save(docx_path)
    logger.info(f"Wrote Word doc: {docx_path}")
    return docx_path

    
if __name__ == "__main__":
    tables = ["mcc_mnc_table", "traforama_isp_list", "mideye_mobile_network_list"]
    report = combined_pipeline(
        user_query="Describe everything comparing india and pakistan",
        sql_tables=tables,
        test_names=["signal"],
        only="",
        horizon=30
    )
    print(report)
'''


# TRUE ASYNC w/ GEMINI AI INvoke -> not sure if I have acess to this?

# final_pipeline.py

import os
import logging
import duckdb
import broadsql
from datacenter import run_scrape_and_markdown
from ooni import scrape_ooni_explorer
from cloudflare import fetch_and_format_markdown
from country_code_converter import get_alpha2_from_country_name
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.prompts import PromptTemplate
import asyncio

from docx import Document
from docx.shared import Pt
from docx2pdf import convert
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# — Init —
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Connect to DuckDB database
DB_PATH = os.path.join(os.path.dirname(__file__), "bryan.db")
con = duckdb.connect(DB_PATH)

# LLM Setup (async-capable)
llm = ChatGoogleGenerativeAI(
    model="gemini-2.5-flash",
    google_api_key=GOOGLE_API_KEY,
    temperature=0
)

# ----------------------------------------
# Async section‐specific callers
# ----------------------------------------
async def answer_sql_section(user_query: str, sql_context: str) -> str:
    prompt = PromptTemplate.from_template(
        """
You are a data analyst. Given these SQL‐extracted tables:

{sql_context}

Answer the user’s question:
"{user_query}"

Focus only on the SQL data. Provide a concise Markdown bullet list of the relevant findings.
"""
    ).format(sql_context=sql_context, user_query=user_query)
    resp = await llm.ainvoke(prompt)           # <<< use async invoke
    return resp.content.strip()

async def answer_dc_section(dc_context: str) -> str:
    prompt = PromptTemplate.from_template(
        """
You are a network infrastructure expert. Given these data‐center entries:

{dc_context}

List each data center with its name, type, and address as a Markdown table.
"""
    ).format(dc_context=dc_context)
    resp = await llm.ainvoke(prompt)
    return resp.content.strip()

async def answer_ooni_section(ooni_context: str) -> str:
    prompt = PromptTemplate.from_template(
        """
You are a network measurement specialist. Here are OONI Explorer results:

{ooni_context}

For each country and test, list anomalies vs accessible counts in bullet points.
"""
    ).format(ooni_context=ooni_context)
    resp = await llm.ainvoke(prompt)
    return resp.content.strip()

async def answer_radar_section(radar_context: str) -> str:
    prompt = PromptTemplate.from_template(
        """
You are a traffic analysis expert. Here is Cloudflare Radar data:

{radar_context}

For each metric, produce a Markdown table showing the top categories and their percentages.
"""
    ).format(radar_context=radar_context)
    resp = await llm.ainvoke(prompt)
    return resp.content.strip()

# ----------------------------------------
# Main pipeline (sync wrapper)
# ----------------------------------------
def combined_pipeline(
    user_query: str,
    sql_tables: list[str],
    test_names: list[str],
    only: str = "",
    horizon: int = 30
) -> str:
    """
    1) SQL‐RAG picks countries_list.
    2) Scrape data centers.
    3) Scrape OONI Explorer per country/test.
    4) Fetch Cloudflare Radar per country.
    5) Kick off four async LLM prompts in parallel.
    """

    # --- Build contexts exactly as before ---
    sql_blocks = []
    countries_list: list[str] = []
    for i, table in enumerate(sql_tables):
        df = con.execute(f"SELECT * FROM {table}").df()
        if i == 0:
            countries_list = broadsql.extract_relevant_rows(df, user_query)
        filtered = broadsql.filter_df(df, "Country", countries_list)
        sql_blocks.append(f"### {table}\n{broadsql.df_to_markdown(filtered)}")
    sql_context = "\n\n".join(sql_blocks) or "No SQL data found."

    dc_md = run_scrape_and_markdown(countries_list)

    ooni_blocks = []
    for test_name in test_names:
        for country in countries_list:
            alpha2 = get_alpha2_from_country_name(country) or ""
            md_table, anomaly_cnt, access_cnt = asyncio.run(
                scrape_ooni_explorer(
                    test_name=test_name,
                    country=alpha2,
                    only=only,
                    horizon=horizon
                )
            )
            ooni_blocks.append(
                f"#### {test_name.title()} – {country}\n"
                f"{md_table}\n"
                f"- Anomalies: {anomaly_cnt}   Accessible: {access_cnt}\n"
            )
    ooni_context = "\n\n".join(ooni_blocks) or "No OONI data found."

    date_range = f"{horizon}d"
    radar_blocks = []
    for country in countries_list:
        alpha2 = get_alpha2_from_country_name(country) or ""
        radar_md = fetch_and_format_markdown(country=alpha2, date_range=date_range)
        radar_blocks.append(f"#### {country} (Radar {date_range})\n{radar_md}")
    radar_context = "\n\n".join(radar_blocks) or "No Radar data found."

    # --- Run the four LLM calls truly in parallel ---
    async def run_sections():
        return await asyncio.gather(
            answer_sql_section(user_query, sql_context),
            answer_dc_section(dc_md),
            answer_ooni_section(ooni_context),
            answer_radar_section(radar_context),
        )

    sql_ans, dc_ans, ooni_ans, radar_ans = asyncio.run(run_sections())

    # --- Stitch final report ---
    report = "\n\n".join([
        "## SQL Data", sql_ans,
        "## Data Centers", dc_ans,
        "## OONI Explorer Results", ooni_ans,
        "## Cloudflare Radar", radar_ans
    ])
    return report

# (Export markdown_to_docx, md_to_pdf, docx_to_pdf unchanged)
def markdown_to_docx(md_text: str, docx_path: str = "report.docx") -> str:
    """Convert markdown to Word with proper headings & bullets."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for line in md_text.splitlines():
        if line.startswith("### "):
            h = doc.add_heading(level=3); h.add_run(line[4:])
        elif line.startswith("## "):
            h = doc.add_heading(level=2); h.add_run(line[3:])
        elif line.startswith("# "):
            h = doc.add_heading(level=1); h.add_run(line[2:])
        elif line.startswith("* "):
            p = doc.add_paragraph(style='List Bullet'); p.add_run(line[2:])
        else:
            doc.add_paragraph(line)
    doc.save(docx_path)
    logger.info(f"Wrote Word doc: {docx_path}")
    return docx_path


if __name__ == "__main__":
    tables = ["mcc_mnc_table", "traforama_isp_list", "mideye_mobile_network_list"]
    report = combined_pipeline(
        user_query="Describe everything comparing india and pakistan",
        sql_tables=tables,
        test_names=["signal"],
        only="",
        horizon=30
    )
    print(report)